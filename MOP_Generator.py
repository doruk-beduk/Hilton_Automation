#!/usr/bin/env python3

import argparse
import re
import os
import sys
from typing import Dict, List, Tuple
from ipaddress import ip_address, ip_network

# python-docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.enum.section import WD_ORIENT

# Third-party
try:
    import pandas as pd
    import PyPDF2
except Exception as e:
    print("Missing required packages. Please run:\n  pip install python-docx PyPDF2 pandas", file=sys.stderr)
    raise

# Base canonical service map
SERVICE_MAP = {
    "https": "443",
    "http": "80",
    "www": "80",
    "ssh": "22",
    "ftp": "21",
    "ftp-data": "20",
    "sqlnet": "1521",
    "tacacs": "49",
    "netbios-ssn": "139",
    "pop3": "110",
}

def _norm_service_key(s: str) -> str:
    # lower-case and remove non [a-z0-9]; "HTTP-s" -> "https", "ftp_data" -> "ftpdata"
    return re.sub(r'[^a-z0-9]+', '', (s or '').lower())

# Rebuild SERVICE_MAP with normalized keys
_SERVICE_MAP_RAW = SERVICE_MAP
SERVICE_MAP = { _norm_service_key(k): v for k, v in _SERVICE_MAP_RAW.items() }

# Extra hints: common service *group* names that imply a single well-known port
# (used when ACL references "object-group <name>" but the group definition isn't present)
SERVICE_GROUP_PORT_HINTS = {
    "https": "443",
    "httpssl": "443",   # "HTTP-SSL" or "HTTPS-SSL" normalized
    "httpss": "443",    # some folks name it oddly
    "https443": "443",
    "ssl": "443",
    "http": "80",
    "www": "80",
    "ssh": "22",
    "ftp": "21",
    "ftpdata": "20",
    "sqlnet": "1521",
    "tacacs": "49",
}

# Always-include rules at the top of Table 6 (Meraki L3 rules)
DEFAULT_PREPEND_RULES = [
    {
        "Policy": "Allow",
        "Description": "",
        "Protocol": "Any",
        "Source": "Qualys-Scanners",
        "Src Port": "Any",
        "Destination": "Any",
        "Dest Port": "Any",
        "Log": "Yes",
    },
    {
        "Policy": "Allow",
        "Description": "",
        "Protocol": "Any",
        "Source": "GISDiscovery",
        "Src Port": "Any",
        "Destination": "Any",
        "Dest Port": "Any",
        "Log": "Yes",
    },
]

def extract_hostname(text: str) -> str:
    m = re.search(r'^\s*hostname\s+(\S+)', text, flags=re.MULTILINE | re.IGNORECASE)
    return m.group(1) if m else "UNKNOWN-HOST"

def set_table_borders(table):
    tbl = table._element
    tblBorders = parse_xml(r"""
        <w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:val="single" w:sz="6" w:space="0" w:color="000000"/>
            <w:left w:val="single" w:sz="6" w:space="0" w:color="000000"/>
            <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>
            <w:right w:val="single" w:sz="6" w:space="0" w:color="000000"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        </w:tblBorders>
    """)
    tbl.tblPr.tblBorders = tblBorders

def set_table_borders(table):
    """
    Apply consistent borders to the entire Word table.
    Adds outer borders and inner grid lines.
    """
    tbl = table._element
    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")      # solid line
        border.set(qn("w:sz"), "6")            # thickness (6 = 0.5pt)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")    # black
        tblBorders.append(border)

    tbl.tblPr.append(tblBorders)

def set_tbl_cell_margins(table, top=30, left=40, bottom=30, right=40):
    """Shrink default cell padding (values are twips; default ~108)."""
    tbl_pr = table._element.tblPr
    tbl_cell_mar = tbl_pr.tblCellMar
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement('w:tblCellMar')
        tbl_pr.append(tbl_cell_mar)
    for side, val in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        el = getattr(tbl_cell_mar, f'w:{side}', None)
        if el is None:
            el = OxmlElement(f'w:{side}')
            tbl_cell_mar.append(el)
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')

def shade_header(row, hexcolor="2F5597"):  # Cisco-ish blue
    for cell in row.cells:
        cell._element.get_or_add_tcPr().append(
            parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hexcolor}"/>')
        )
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = None  # keep Word default (usually white on dark bg)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def standardize_cells(table, font_size_pt=10):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                for run in p.runs:
                    run.font.size = Pt(font_size_pt)

def set_no_wrap(cell):
    """Prevent Word from wrapping text inside this cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(OxmlElement('w:noWrap'))

def nobreak_text(s: str) -> str:
    """Replace spaces and hyphens with non-breaking versions to avoid mid-word wraps."""
    if s is None:
        return ""
    return s.replace(' ', '\u00A0').replace('-', '\u2011')  # NBSP and NB hyphen

def read_config_text(path: str) -> str:
    """
    Returns a normalized text version of the ASA config.
    Prefers .txt; supports .pdf via PyPDF2 text extraction.
    """
    if not os.path.exists(path):
        raise FileNotFoundError(path)

    text = ""
    if path.lower().endswith(".pdf"):
        reader = PyPDF2.PdfReader(path)
        for page in reader.pages:
            t = page.extract_text() or ""
            text += t + "\n"
    else:
        with open(path, "r", errors="ignore") as f:
            text = f.read()

    # Normalize newlines
    text = text.replace("\r", "\n")
    text = re.sub(r'# Cisco Confidential.*', '', text)

    # Split long PDF lines but preserve true line breaks
    # Do NOT collapse all whitespace — that breaks block detection
    text = re.sub(r' +', ' ', text)  # collapse multiple spaces, not newlines

    # Ensure key tokens start on their own line (helps if PDF extractor merges)
    tokens = ["access-list ", "object-group network", "object network"]
    for token in tokens:
        text = re.sub(rf' *{token}', f"\n{token}", text)

    # Clean up multiple newlines
    text = re.sub(r'\n+', '\n', text)

    lines = [l.strip() for l in text.split("\n") if l.strip()]
    return "\n".join(lines)


def bits_from_mask(mask: str) -> int:
    return sum(bin(int(o)).count("1") for o in mask.split("."))


def parse_objects(text: str) -> Dict[str, Dict[str, str]]:
    """
    Parse 'object network <name> host <ip>'  OR  'object network <name> <ip> <mask>' (subnet-inlined) 
    and also the block style 'object network <name>' followed by 'host ...' or 'subnet ...'.
    Returns: { name: {"name": name, "cidr": "<ip/prefix>"} }
    """
    objects = {}

    # Single-line (after normalization often becomes single-line)
    for m in re.finditer(r'\bobject network ([A-Za-z0-9_\-\.]+)\s+(host\s+(\d+\.\d+\.\d+\.\d+)|(\d+\.\d+\.\d+\.\d+)\s+(\d+\.\d+\.\d+\.\d+))\b', text):
        name = m.group(1)
        if m.group(3):  # host X.X.X.X
            objects[name] = {"name": name, "cidr": f"{m.group(3)}/32"}
        else:  # X.X.X.X Y.Y.Y.Y
            ip, mask = m.group(4), m.group(5)
            objects[name] = {"name": name, "cidr": f"{ip}/{bits_from_mask(mask)}"}

    # Block-style fallback: object network <name>\n (host|subnet) ...
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        m = re.match(r'object network ([A-Za-z0-9_\-\.]+)', lines[i])
        if m:
            name = m.group(1)
            j = i + 1
            # Look only at the immediate next line(s) until a new block starts
            while j < len(lines) and not re.match(r'(object|object-group|access-list)\b', lines[j]):
                mh = re.search(r'\bhost\s+(\d+\.\d+\.\d+\.\d+)\b', lines[j])
                ms = re.search(r'\bsubnet\s+(\d+\.\d+\.\d+\.\d+)\s+(\d+\.\d+\.\d+\.\d+)\b', lines[j])
                if mh:
                    objects[name] = {"name": name, "cidr": f"{mh.group(1)}/32"}
                    break
                if ms:
                    ip, mask = ms.group(1), ms.group(2)
                    objects[name] = {"name": name, "cidr": f"{ip}/{bits_from_mask(mask)}"}
                    break
                j += 1
            i = j
            continue
        i += 1

    return objects

def parse_service_groups(text: str) -> Dict[str, List[Dict[str, str]]]:
    """
    Parse:
      object-group service <NAME> [<proto>]
        port-object eq <p>
        port-object range <a> <b>
        service-object <proto> [destination] eq|range <...>
        service-object icmp [type]  (treated as protocol=ICMP, ports Any)

    Returns:
      { group_name: [ { "protocol": "TCP|UDP|ICMP|ANY",
                        "type": "port|range|any",
                        "port": "80" | "",
                        "portEnd": "443" | "" } , ... ] }
    """
    svc_groups: Dict[str, List[Dict[str, str]]] = {}
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        mg = re.match(r'object-group\s+service\s+([A-Za-z0-9_\-\.]+)(?:\s+(\w+))?', lines[i])
        if not mg:
            i += 1
            continue

        gname = mg.group(1)
        default_proto = (mg.group(2) or "").upper()  # may be blank; entries can carry proto
        i += 1
        members: List[Dict[str, str]] = []

        while i < len(lines) and not re.match(r'(object-group\s+|object\s+|access-list\s+)', lines[i]):
            line = lines[i].strip()

            # port-object eq 443
            m_eq = re.match(r'port-object\s+eq\s+(\d+)', line, flags=re.I)
            if m_eq:
                members.append({"protocol": default_proto or "ANY", "type": "port", "port": m_eq.group(1), "portEnd": ""})
                i += 1
                continue

            # port-object range 8000 8080
            m_rg = re.match(r'port-object\s+range\s+(\d+)\s+(\d+)', line, flags=re.I)
            if m_rg:
                members.append({"protocol": default_proto or "ANY", "type": "range", "port": m_rg.group(1), "portEnd": m_rg.group(2)})
                i += 1
                continue

            # service-object tcp destination eq 443
            m_srv = re.match(r'service-object\s+(\w+)(?:\s+destination)?\s+(eq|range)\s+(\S+)(?:\s+(\S+))?', line, flags=re.I)
            if m_srv:
                proto = m_srv.group(1).upper()
                kind = m_srv.group(2).lower()
                p1 = m_srv.group(3)
                p2 = m_srv.group(4) or ""
                if kind == "eq":
                    members.append({"protocol": proto, "type": "port", "port": service_to_port(p1), "portEnd": ""})
                else:
                    members.append({"protocol": proto, "type": "range", "port": service_to_port(p1), "portEnd": service_to_port(p2)})
                i += 1
                continue

            # service-object icmp (optionally with type) → treat as protocol ICMP
            m_icmp = re.match(r'service-object\s+icmp\b', line, flags=re.I)
            if m_icmp:
                members.append({"protocol": "ICMP", "type": "any", "port": "", "portEnd": ""})
                i += 1
                continue

            i += 1

        svc_groups[gname] = members

    return svc_groups


def referenced_service_groups_from_acls(acl_lines: List[str],
                                        service_groups_defined: Dict[str, List[Dict[str, str]]]) -> List[str]:
    """
    Collect object-group names that appear in ACLs and are defined as service groups.
    """
    refs = set()
    for l in acl_lines:
        for m in re.finditer(r'\bobject-group\s+([A-Za-z0-9_\-\.]+)\b', l):
            name = m.group(1)
            if name in service_groups_defined:
                refs.add(name)
    return sorted(refs)


def parse_object_groups(text: str) -> Dict[str, List[str]]:
    """
    Parse 'object-group network <name>' blocks and collect members as CIDRs.
    Returns: { group_name: ["ip/prefix", ...] }
    """
    groups = {}
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        mg = re.match(r'object-group network ([A-Za-z0-9_\-\.]+)', line)
        if mg:
            gname = mg.group(1)
            i += 1
            members: List[str] = []
            while i < len(lines) and not (
                lines[i].startswith("object-group") or
                lines[i].startswith("object ") or
                lines[i].startswith("access-list ")
            ):
                # "network-object host X.X.X.X"
                m1 = re.search(r'network-object host (\d+\.\d+\.\d+\.\d+)', lines[i])
                # "network-object X.X.X.X Y.Y.Y.Y"
                m2 = re.search(r'network-object (\d+\.\d+\.\d+\.\d+)\s+(\d+\.\d+\.\d+\.\d+)', lines[i])
                if m1:
                    members.append(m1.group(1) + "/32")
                elif m2:
                    ip, mask = m2.group(1), m2.group(2)
                    members.append(f"{ip}/{bits_from_mask(mask)}")
                i += 1
            groups[gname] = members
            continue
        i += 1
    return groups


def service_to_port(token: str) -> str:
    key = _norm_service_key(token)
    return SERVICE_MAP.get(key, token)


def parse_acls(text: str,
               objects: Dict[str, Dict[str, str]],
               service_groups_all: Dict[str, List[Dict[str, str]]]
               ) -> List[Dict[str, str]]:
    """
    Build Meraki L3 rules from ASA ACLs.

    NEW:
    - Handles ACLs that begin with "object-group <service_group>" (no explicit proto).
    - Emits ONE rule per protocol in that service group (TCP, UDP, ICMP).
      * TCP/UDP rows aggregate all ports of that protocol into a comma-separated list.
      * Ranges appear as "a-b".
    - Keeps object names in Source/Destination (no IP expansion).
    """
    def _format_ports(ports: List[Tuple[str, str]]) -> str:
        """
        ports: list of (type, value) where type in {"port","range","any"}.
        For 'port', value is "80"; for 'range', value is "8000-8080"; for 'any', value is "Any".
        Return a user-friendly comma-separated string.
        """
        # if any member is 'any' → the protocol row is Any
        if any(t == "any" for t, _ in ports):
            return "Any"

        singles = []
        ranges  = []
        for t, v in ports:
            if t == "port":
                singles.append(v)
            elif t == "range":
                ranges.append(v)
        # numeric sort singles where possible; keep strings otherwise
        def _key(x):
            try:
                return (0, int(x))
            except ValueError:
                return (1, x)
        singles = sorted(set(singles), key=_key)
        ranges  = sorted(set(ranges), key=lambda s: tuple(int(p) if p.isdigit() else p for p in s.split("-")))
        pieces = singles + ranges
        return ", ".join(pieces) if pieces else "Any"

    def _gather_src_dst(rest: str) -> Tuple[str, str]:
        """Extract Source and Destination tokens keeping object/object-group names literal."""
        occurrences: List[Tuple[int, str, str]] = []
        for kind, pat in [
            ("host", r'host (\d+\.\d+\.\d+\.\d+)'),
            ("object", r'\bobject ([A-Za-z0-9_\-\.]+)\b'),
            ("object-group", r'\bobject-group ([A-Za-z0-9_\-\.]+)\b'),
            ("subnet", r'(\d+\.\d+\.\d+\.\d+)\s+(\d+\.\d+\.\d+\.\d+)')
        ]:
            for mm in re.finditer(pat, rest):
                if kind == "subnet":
                    ip, mask = mm.group(1), mm.group(2)
                    occurrences.append((mm.start(), f"{ip}/{bits_from_mask(mask)}", kind))
                else:
                    occurrences.append((mm.start(), mm.group(1), kind))
        occurrences.sort(key=lambda x: x[0])

        src = "Any"
        dst = "Any"
        if occurrences:
            v, k = occurrences[0][1], occurrences[0][2]
            src = v  # keep names literal (no IP expansion)
        if len(occurrences) >= 2:
            v, k = occurrences[1][1], occurrences[1][2]
            dst = v
        return src, dst

    lines = text.splitlines()
    rules: List[Dict[str, str]] = []
    last_remark = ""

    for i, line in enumerate(lines):
        # Track description from remark
        mrem = re.match(r'access-list \S+ remark (.+)', line)
        if mrem:
            last_remark = mrem.group(1).strip()
            continue

        # --- CASE A: ACL uses explicit protocol (existing behavior) ---
        m = re.match(r'access-list (\S+) extended (permit|deny) (\w+)\s+(.+)', line)
        # But be careful: if (\w+) actually captured "object-group", we should treat as CASE B
        if m and m.group(3).lower() != "object":
            acl_name, action, proto = m.group(1), m.group(2), m.group(3)
            rest = m.group(4)
            # If this was actually "object-group <svc>", (\w+) will be "object-group" → handle below
            if proto.lower() != "object-group":
                proto_up = "Any" if proto.lower() == "ip" else proto.upper()

                # Destination port:
                dst_port = "Any"
                meq = re.search(r'\beq\s+([A-Za-z0-9_\-\/]+)', rest)
                if meq:
                    dst_port = service_to_port(meq.group(1))

                if dst_port == "Any" and proto_up in ("TCP", "UDP"):
                    # Check if a service object-group is referenced on the ACL line
                    svc_group_names = [mm.group(1) for mm in re.finditer(r'\bobject-group\s+([A-Za-z0-9_\-\.]+)\b', rest)]
                    for gname in svc_group_names:
                        members = service_groups_all.get(gname, [])
                        if not members:
                            # name-based hint fallback
                            hint = SERVICE_GROUP_PORT_HINTS.get(_norm_service_key(gname))
                            if hint:
                                dst_port = hint
                                break
                            continue
                        # gather all port items for the current proto
                        collected: List[Tuple[str, str]] = []
                        for mentry in members:
                            if (mentry.get("protocol") or "").upper() not in (proto_up, "ANY"):
                                continue
                            typ = (mentry.get("type") or "").lower()
                            p1  = (mentry.get("port") or "").strip()
                            p2  = (mentry.get("portEnd") or "").strip()
                            if typ == "any":
                                collected.append(("any", "Any"))
                            elif typ == "port" and p1:
                                collected.append(("port", service_to_port(p1)))
                            elif typ == "range" and p1 and p2:
                                collected.append(("range", f"{service_to_port(p1)}-{service_to_port(p2)}"))
                        if collected:
                            dst_port = _format_ports(collected)
                            break

                src, dst = _gather_src_dst(rest)
                src_port = "Any"
                if proto_up == "ICMP":
                    dst_port = "Any"

                rules.append({
                    "Policy": "Allow" if action == "permit" else "Deny",
                    "Description": last_remark,
                    "Protocol": proto_up,
                    "Source": src,
                    "Src Port": src_port,
                    "Destination": dst,
                    "Dest Port": dst_port,
                    "Log": "Yes"
                })
                last_remark = ""
                continue  # next line

        # --- CASE B: ACL begins with "object-group <service_group>" (no explicit proto) ---
        msvc = re.match(r'access-list (\S+) extended (permit|deny)\s+object-group\s+([A-Za-z0-9_\-\.]+)\s+(.+)', line)
        if not msvc:
            continue

        acl_name, action, svc_group_name, rest = msvc.groups()
        members = service_groups_all.get(svc_group_name, [])

        # group ports by protocol
        by_proto: Dict[str, List[Tuple[str, str]]] = {"TCP": [], "UDP": [], "ICMP": []}
        for mentry in members:
            proto_m = (mentry.get("protocol") or "").upper()
            typ     = (mentry.get("type") or "any").lower()
            p1      = (mentry.get("port") or "").strip()
            p2      = (mentry.get("portEnd") or "").strip()

            # If a member has no proto (e.g., group header specified), treat as ANY → add to all relevant?
            if proto_m in ("", "ANY"):
                targets = ("TCP", "UDP")
                # If it's an ICMP member, typ may be 'any' and proto_m might be 'ICMP'
            else:
                targets = (proto_m,)

            # ICMP rows don't have ports
            if "ICMP" in targets or proto_m == "ICMP":
                by_proto.setdefault("ICMP", [])
                by_proto["ICMP"].append(("any", "Any"))
                continue

            # TCP/UDP ports
            for tgt in targets:
                if typ == "any":
                    by_proto.setdefault(tgt, [])
                    by_proto[tgt].append(("any", "Any"))
                elif typ == "port" and p1:
                    by_proto.setdefault(tgt, [])
                    by_proto[tgt].append(("port", service_to_port(p1)))
                elif typ == "range" and p1 and p2:
                    by_proto.setdefault(tgt, [])
                    by_proto[tgt].append(("range", f"{service_to_port(p1)}-{service_to_port(p2)}"))

        src, dst = _gather_src_dst(rest)

        # one rule per protocol with aggregated ports
        for proto_up in ("TCP", "UDP", "ICMP"):
            items = by_proto.get(proto_up, [])
            if not items:
                continue
            dst_port = "Any" if proto_up == "ICMP" else _format_ports(items)

            rules.append({
                "Policy": "Allow" if action == "permit" else "Deny",
                "Description": last_remark,
                "Protocol": proto_up if proto_up != "ICMP" else "ICMP",
                "Source": src,
                "Src Port": "Any",
                "Destination": dst,
                "Dest Port": dst_port,
                "Log": "Yes"
            })
        last_remark = ""

    return rules


def extract_acl_lines(text: str) -> List[str]:
    """Only ACL lines that start with 'access-list' (allow leading spaces) and are NOT remarks."""
    lines = text.splitlines()
    out = []
    for l in lines:
        if re.match(r'^\s*access-list\s+', l) and not re.search(r'\bremark\b', l):
            out.append(l.strip())
    return out

def referenced_objects_from_acls(acl_lines: List[str],
                                 objects_defined: Dict[str, Dict[str, str]]) -> List[str]:
    """
    From ACL lines, collect 'object <name>' references that actually exist as
    defined network objects (host/subnet) and return the deduped list of names.
    """
    refs = set()
    for l in acl_lines:
        for m in re.finditer(r'\bobject\s+([A-Za-z0-9_\-\.]+)\b', l):
            name = m.group(1)
            if name in objects_defined:
                refs.add(name)
    return sorted(refs)

def referenced_network_groups_from_acls(acl_lines: List[str],
                                        network_groups_defined: Dict[str, List[str]]) -> List[str]:
    """
    From ACL lines, collect object-group references and keep only those that exist
    as *network* object-groups (i.e., in parse_object_groups()).
    """
    refs = set()
    for l in acl_lines:
        for m in re.finditer(r'\bobject-group\s+([A-Za-z0-9_\-\.]+)\b', l):
            name = m.group(1)
            if name in network_groups_defined:  # only network object-groups
                refs.add(name)
    return sorted(refs)



def build_tables_from_config_text(text: str):
    """
    Build the three tables for the MoP:
      - Table 4: Meraki Object Groups (ACL-referenced network object-groups)
      - Table 5: Meraki Objects (ACL-referenced local objects + members of Table-4 groups + members of referenced service groups)
      - Table 6: L3 Meraki Firewall Rules (from ACLs) with two standard 'always-allow' rows prepended
    """
    # ---- Parse primitives from ASA text ----
    objects = parse_objects(text)                      # { name: {"cidr": "<ip/prefix>"} }
    groups_all = parse_object_groups(text)             # { group_name: [cidr_list] }
    service_groups_all = parse_service_groups(text)    # { svc_group: [ {protocol,type,port,portEnd}, ... ] }
    acl_only = extract_acl_lines(text)                 # list[str] of non-remark ACL lines

    # ---- Table 4: only network groups referenced in ACLs ----
    ref_groups = referenced_network_groups_from_acls(acl_only, groups_all)
    df_groups = pd.DataFrame([{"name": g, "category": "NetworkObjectGroup"} for g in ref_groups])

    # ---- Table 6: rules from ACLs ----
    rules = parse_acls(
        "\n".join(acl_only),
        {k: {"ip": v["cidr"].split("/")[0]} for k, v in objects.items()},
        service_groups_all,
    )


    _OBJ_IP_NAME = re.compile(
        r'^(?:obj|object|host|addr|ip|net)[-_]?((?:\d{1,3}\.){3}\d{1,3})(?:/\d{1,2})?$', re.I
    )
    _INLINE_IP = re.compile(r'(?<!\d)(?:\d{1,3}\.){3}\d{1,3}(?!\d)')

    def _squash_obj_ip(token: str) -> str:
        s = (token or "").strip()
        if not s:
            return s
        m = _OBJ_IP_NAME.match(s)
        if m:
            return m.group(1)
        m2 = _INLINE_IP.search(s)
        if m2:
            return m2.group(0)
        return s

    for r in rules:
        r["Source"] = _squash_obj_ip(r.get("Source", ""))
        r["Destination"] = _squash_obj_ip(r.get("Destination", ""))

    # Canonical names we always prepend
    DEFAULT_PREPEND_RULES = [
        {"Policy":"Allow","Description":"","Protocol":"Any","Source":"Qualys-Scanners","Src Port":"Any","Destination":"Any","Dest Port":"Any","Log":"Yes"},
        {"Policy":"Allow","Description":"","Protocol":"Any","Source":"GISDiscovery","Src Port":"Any","Destination":"Any","Dest Port":"Any","Log":"Yes"},
    ]

    # Aliases → canonical
    SUPPRESS_SOURCE_ALIASES = {
        "qualysscansrvrs.hilton.com":"qualys-scanners",
        "qualys_scanners":"qualys-scanners",
        "qualys scanners":"qualys-scanners",
        "qualys-scanners":"qualys-scanners",
        "gis_discovery":"gisdiscovery",
        "gis discovery":"gisdiscovery",
        "gisdiscovery":"gisdiscovery",
    }
    ALWAYS_ALLOW_SOURCES = {"qualys-scanners","gisdiscovery"}

    def _canon_source(s: str) -> str:
        k = (s or "").strip().lower().replace(" ", "").replace("_", "")
        return SUPPRESS_SOURCE_ALIASES.get(k, k)

    # Collect CIDRs backing the always-allow names (if they exist in the config)
    always_cidrs = []
    for nm in ("Qualys-Scanners","GISDiscovery"):
        if nm in objects:            # object network
            always_cidrs.append(objects[nm]["cidr"])
        if nm in groups_all:         # object-group network
            always_cidrs.extend(groups_all[nm])

    def _source_in_always_sets(src: str) -> bool:
        s = (src or "").strip()
        # 1) name match
        if _canon_source(s) in ALWAYS_ALLOW_SOURCES:
            return True
        # 2) IP / CIDR inside any of the always-allow CIDRs
        try:
            if "/" in s:
                net = ip_network(s, strict=False)
                return any(net.subnet_of(ip_network(c, strict=False)) for c in always_cidrs)
            else:
                ip = ip_address(s)
                return any(ip in ip_network(c, strict=False) for c in always_cidrs)
        except ValueError:
            return False

    def _is_generic_always_allow(row: dict) -> bool:
        return (
            (row.get("Policy") or "").lower() == "allow" and
            (row.get("Protocol") or "").lower() == "any" and
            (row.get("Src Port") or "").lower() == "any" and
            (row.get("Dest Port") or "").lower() == "any" and
            (row.get("Destination") or "").lower() in {"any","any/0"} and
            _source_in_always_sets(row.get("Source",""))
        )

    # Drop ACL-derived duplicates of the default rows
    rules = [r for r in rules if not _is_generic_always_allow(r)]

    # Now prepend your two standard rows
    rules = DEFAULT_PREPEND_RULES + rules

    # RIght here
    # --- Order Deny rules last + append terminal local-subnet denies ---

    def _looks_ip_or_cidr(s: str) -> bool:
        s = (s or "").strip()
        if not s: return False
        try:
            if "/" in s:
                ip_network(s, strict=False); return True
            ip_address(s); return True
        except ValueError:
            return False

    def _cidr24_from_host(ip_str: str) -> str:
        ip = ip_address(ip_str)
        # /24 for IPv4 (assumption matches Hilton pattern)
        return f"{ip.exploded.rsplit('.', 1)[0]}.0/24"

    def _collect_candidate_subnets(rules_list):
        """Find best local /24 to deny at end; prefer explicit subnets, else infer from hosts.
        Prioritize 10.122.x.x (if present), else 192.168.x.x."""
        explicit = []   # explicit CIDRs seen anywhere in rules
        inferred = []   # /24s inferred from hosts
        re_host = re.compile(r'^(?:\d{1,3}\.){3}\d{1,3}$')
        for r in rules_list:
            for field in ("Source","Destination"):
                tok = (r.get(field) or "").strip()
                if not tok: continue
                # explicit CIDR?
                if "/" in tok:
                    try: explicit.append(str(ip_network(tok, strict=False)))
                    except ValueError: pass
                # lone host IP?
                elif re_host.match(tok):
                    try:
                        ip = ip_address(tok)
                        # focus on RFC1918; especially 10.122.* and 192.168.*
                        if str(ip).startswith("10.122."):
                            inferred.append(_cidr24_from_host(tok))
                        elif str(ip).startswith("192.168."):
                            inferred.append(_cidr24_from_host(tok))
                        elif str(ip).startswith("10."):
                            # de-prioritize generic 10/8 unless nothing else found
                            inferred.append(_cidr24_from_host(tok))
                    except ValueError:
                        pass

        # Prefer explicit 10.122.*/* then explicit 192.168.*/*
        explicit_10122 = [c for c in explicit if c.startswith("10.122.")]
        explicit_192168 = [c for c in explicit if c.startswith("192.168.")]
        if explicit_10122: return explicit_10122[0]
        if explicit_192168: return explicit_192168[0]

        # Else pick the most frequent inferred /24, with 10.122.* favored
        if inferred:
            from collections import Counter
            counts = Counter(inferred)
            # best by count, prefer 10.122.* on tie
            best = max(counts.items(), key=lambda kv: (kv[1], kv[0].startswith("10.122.")))
            return best[0]

        return ""  # nothing found

    def _same_row(a: dict, b: dict) -> bool:
        # Compare the visible Meraki columns only
        keys = ("Policy","Description","Protocol","Source","Src Port","Destination","Dest Port","Log")
        return all((a.get(k) or "").strip().lower() == (b.get(k) or "").strip().lower() for k in keys)

    def _ensure_terminal_local_denies(rules_list):
        local = _collect_candidate_subnets(rules_list)
        if not local:
            return rules_list  # nothing to add

        deny_src = {
            "Policy":"Deny","Description":"","Protocol":"Any",
            "Source":local,"Src Port":"Any","Destination":"Any","Dest Port":"Any","Log":"Yes"
        }
        deny_dst = {
            "Policy":"Deny","Description":"","Protocol":"Any",
            "Source":"Any","Src Port":"Any","Destination":local,"Dest Port":"Any","Log":"Yes"
        }

        # Avoid duplicates if ACEs already include them
        have_src = any(_same_row(r, deny_src) for r in rules_list)
        have_dst = any(_same_row(r, deny_dst) for r in rules_list)

        # Move all existing Deny rules to the end (keep relative order)
        allows = [r for r in rules_list if (r.get("Policy") or "").lower() == "allow"]
        denies = [r for r in rules_list if (r.get("Policy") or "").lower() != "allow"]

        if not have_src: denies.append(deny_src)
        if not have_dst: denies.append(deny_dst)

        return allows + denies

    # Reorder + append the two terminal local-subnet Deny rules
    rules = _ensure_terminal_local_denies(rules)

    # --- Drop redundant blanket denies + de-duplicate identical rows ---

    def _is_blanket_deny(r: dict) -> bool:
        return (
            (r.get("Policy") or "").strip().lower() == "deny" and
            (r.get("Protocol") or "").strip().lower() in {"any", "ip", ""} and
            (r.get("Source") or "").strip().lower() in {"any", "any/0"} and
            (r.get("Destination") or "").strip().lower() in {"any", "any/0"} and
            (r.get("Src Port") or "").strip().lower() in {"any", ""} and
            (r.get("Dest Port") or "").strip().lower() in {"any", ""} and
            (r.get("Description") or "").strip() == ""
        )

    # Remove ACL-derived blanket deny rows (Meraki has implicit deny and you have explicit local denies)
    rules = [r for r in rules if not _is_blanket_deny(r)]

    # Optional: if you ever want to keep exactly one blanket deny, replace the line above with:
    # keeps_one = True
    # seen_blanket = False
    # filtered = []
    # for r in rules:
    #     if _is_blanket_deny(r):
    #         if keeps_one and not seen_blanket:
    #             filtered.append(r); seen_blanket = True
    #         # else drop it
    #     else:
    #         filtered.append(r)
    # rules = filtered

    # De-dupe any fully identical rows that may have come from the ACLs & our appends
    def _row_key(r: dict) -> tuple:
        cols = ("Policy","Description","Protocol","Source","Src Port","Destination","Dest Port","Log")
        return tuple((r.get(c) or "").strip().lower() for c in cols)

    _seen = set()
    deduped = []
    for r in rules:
        k = _row_key(r)
        if k in _seen:
            continue
        _seen.add(k)
        deduped.append(r)
    rules = deduped

    # ------------------------------------------------------------------
    df_rules = pd.DataFrame(rules)
    df_rules.insert(0, "#", range(1, len(df_rules) + 1))

    # ---- Table 5: objects to materialize ----
    obj_rows = []

    # A) ACL-referenced local 'object network <name>'
    ref_objs = referenced_objects_from_acls(acl_only, objects)
    for name in ref_objs:
        cidr = objects[name]["cidr"]
        obj_rows.append({
            "name": name,
            "category": "network",
            "type": "cidr",
            "cidr": cidr,
            "fqdn": "None",
            "groupIDs": ""
        })

    # B) Expand ONLY the network groups shown in Table 4
    for gname in ref_groups:
        for cidr in groups_all.get(gname, []):
            ip_part = cidr.split("/")[0]
            safe = ip_part.replace(".", "-")
            obj_rows.append({
                "name": f"{gname}_host_{safe}",
                "category": "network",
                "type": "cidr",
                "cidr": cidr if "/" in cidr else f"{cidr}/32",
                "fqdn": "None",
                "groupIDs": gname
            })

    # C) Expand service object-groups referenced by ACLs (summarize as 'service' rows)
    ref_svc_groups = referenced_service_groups_from_acls(acl_only, service_groups_all)
    for sg in ref_svc_groups:
        for m in service_groups_all.get(sg, []):
            typ = (m.get("type") or "any").lower()
            p1 = (m.get("port") or "").strip()
            p2 = (m.get("portEnd") or "").strip()

            if typ == "range" and p1 and p2:
                member_name = f"{sg}_range_{p1}-{p2}"
            elif typ == "port" and p1:
                member_name = f"{sg}_port_{p1}"
            else:
                member_name = f"{sg}_any"

            obj_rows.append({
                "name": member_name,
                "category": "service",
                "type": typ,
                "cidr": "",
                "fqdn": "None",
                "groupIDs": sg
            })

    df_objects = (
        pd.DataFrame(obj_rows)
        .drop_duplicates(subset=["name", "category", "type", "cidr", "fqdn", "groupIDs"])
    )

    return df_groups, df_objects, df_rules



def write_mop_docx(df_groups, df_objects, df_rules, out_path: str, hostname: str = "UNKNOWN-HOST"):
    doc = Document()
    
    # Define custom heading font sizes
    style_h1 = doc.styles['Heading 1']
    style_h1.font.size = Pt(18)
    style_h1.font.bold = True

    style_h2 = doc.styles['Heading 2']
    style_h2.font.size = Pt(16)
    style_h2.font.bold = True

    style_h3 = doc.styles['Heading 3']
    style_h3.font.size = Pt(14)
    style_h3.font.bold = True

    p = doc.add_heading('4. Method of Procedures', level=1)

    run = p.runs[0]
    run.font.size = Pt(18)      
    doc.add_paragraph(
        "This Method of Procedure (MoP) represents the sequential steps to complete the migration of the firewall rules."
    )

    # ---------- 4.1.1 Migration Plan ----------
    doc.add_heading(f'4.1.1  Migration Plan – {hostname}', level=3)

    # Numbered items
    p1 = doc.add_paragraph("1. Review existing ASA configuration")
    p1.style = 'List Number'

    p2 = doc.add_paragraph("2. Create Policy Objects for the firewall rules")
    p2.style = 'List Number'

    # Lettered sub-steps 
    from docx.shared import Inches
    sub = [
        "a. Navigate to Organization -> Configure -> Policy Objects",
        "b. Add new",
        "c. Add the following Objects and Groups",
    ]
    for s in sub:
        sp = doc.add_paragraph(s)
        sp.paragraph_format.left_indent = Inches(0.5)

    # Add the extra-indented note
    note = doc.add_paragraph("a. Note: Objects and Object Groups will be imported via Meraki scripting")
    note.paragraph_format.left_indent = Inches(1.0)   # ← one level deeper

    for _ in range(1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # -------- Table 4 --------
    doc.add_heading('Table 4: Meraki Object Groups', level=2)

    hdr4 = ['name', 'category']
    t4 = doc.add_table(rows=1, cols=len(hdr4))

    # Match Table 5 & 6 behavior
    t4.autofit = True

    # Starter widths that keep columns proportional on landscape page
    col_in_4 = [5, 5]  # ~9" total usable width with 0.6" margins
    for i, w in enumerate(col_in_4):
        t4.columns[i].width = Inches(w)

    # Header
    for i, h in enumerate(hdr4):
        cell = t4.rows[0].cells[i]
        cell.text = h
        set_no_wrap(cell)
    shade_header(t4.rows[0])

    # Rows
    for _, r in df_groups.iterrows():
        cells = t4.add_row().cells
        vals = [str(r.get(h, "")) for h in hdr4]

        vals[0] = nobreak_text(vals[0])
        vals[1] = nobreak_text(vals[1])

        for i, v in enumerate(vals):
            cells[i].text = v
            set_no_wrap(cells[i])

    set_table_borders(t4)
    standardize_cells(t4, font_size_pt=9)

    for _ in range(2):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # -------- Table 5 --------
    # Switch to landscape orientation
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE

    # Swap page width and height to match orientation
    section.page_width, section.page_height = section.page_height, section.page_width

    # Optional: adjust margins for extra space
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    
    doc.add_heading('Table 5: Meraki Objects', level=2)

    hdr5 = ['name', 'category', 'type', 'cidr', 'fqdn', 'groupIDs']
    t5 = doc.add_table(rows=1, cols=len(hdr5))
    t5.autofit = True

    # Starter widths that fit; Word can still shrink as needed
    col_in = [2.1, 1.4, 1.3, 1.8, 1.3, 2.1]  # total ~10.0 in (fits US Letter with margins)
    for i, w in enumerate(col_in):
        t5.columns[i].width = Inches(w)

    for i, h in enumerate(hdr5):
        t5.rows[0].cells[i].text = h
    shade_header(t5.rows[0])

    for _, r in df_objects.iterrows():
        cells = t5.add_row().cells
        vals = [str(r.get(h, "")) for h in hdr5]

        # Non-breaking for common offenders
        vals[0] = nobreak_text(vals[0])  # name
        vals[3] = nobreak_text(vals[3])  # cidr
        vals[5] = nobreak_text(vals[5])  # groupIDs

        for i, v in enumerate(vals):
            cells[i].text = v

        # Keep key columns from wrapping badly
        for idx in (0, 3, 5):
            set_no_wrap(cells[idx])

    set_table_borders(t5)
    standardize_cells(t5, font_size_pt=9)  # slightly smaller font to fit better

    for _ in range(2):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)


    # -------- Table 6 --------
    doc.add_heading('Table 6: L3 Meraki Firewall Rules', level=2)

    hdr6 = ['#', 'Policy', 'Description', 'Protocol', 'Source', 'Src Port', 'Destination', 'Dest Port', 'Log']
    t6 = doc.add_table(rows=1, cols=len(hdr6))

    # Match Table 5 behavior: give sane starting widths and let Word auto-fit
    t6.autofit = True

    # Starter widths (inches) that fit on a landscape page with ~0.6" margins
    # Bias more space to Description, Source, Destination
    col_in_6 = [0.5, 0.9, 2.6, 0.9, 2.2, 0.9, 2.2, 1.0, 0.6]
    for i, w in enumerate(col_in_6):
        t6.columns[i].width = Inches(w)

    # Header
    for i, h in enumerate(hdr6):
        cell = t6.rows[0].cells[i]
        cell.text = h
        set_no_wrap(cell)  # keep headers on one line
    shade_header(t6.rows[0])

    # Rows
    for _, r in df_rules.iterrows():
        cells = t6.add_row().cells
        for i, h in enumerate(hdr6):
            val = str(r[h])

            # Prevent ugly mid-token breaks on key fields
            if h in ('Source', 'Destination', 'Dest Port', 'Src Port', 'Protocol'):
                val = nobreak_text(val)
                set_no_wrap(cells[i])

            cells[i].text = val

    set_table_borders(t6)
    standardize_cells(t6, font_size_pt=9)

    doc.save(out_path)


def generate_mop(input_path: str, output_docx: str):
    text = read_config_text(input_path)
    hostname = extract_hostname(text)

    df_groups, df_objects, df_rules = build_tables_from_config_text(text)

    base = os.path.splitext(output_docx)[0]
    df_groups.to_csv(base + "_table4_groups.csv", index=False)
    df_objects.to_csv(base + "_table5_objects.csv", index=False)
    df_rules.to_csv(base + "_table6_rules.csv", index=False)

    write_mop_docx(df_groups, df_objects, df_rules, output_docx, hostname=hostname)


def main():
    ap = argparse.ArgumentParser(description="Generate Meraki MoP tables from an ASA config (TXT preferred; PDF supported).")
    ap.add_argument("--input", "-i", required=True, help="Path to ASA config file (.txt preferred, .pdf supported)")
    ap.add_argument("--output", "-o", required=True, help="Path to output Word document (.docx)")
    args = ap.parse_args()

    generate_mop(args.input, args.output)
    print(f"Generated MoP: {args.output}")
    base = os.path.splitext(args.output)[0]
    print(f"Audit CSVs: {base}_table4_groups.csv, {base}_table5_objects.csv, {base}_table6_rules.csv")


if __name__ == "__main__":
    main()
